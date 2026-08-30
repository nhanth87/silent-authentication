#!/usr/bin/env python3
"""Assemble Restlink Silent Auth proposal DOCX (~50 pages) from markdown chapters + charts."""

from __future__ import annotations

import json
import re
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
CHAPTERS = ROOT / "chapters"
ASSETS = ROOT / "assets"
CHARTS = ASSETS / "charts"
OUT = ROOT / "Restlink_Silent_Auth_Proposal_v3.docx"

GREEN = RGBColor(0x07, 0x89, 0x30)
DARK = RGBColor(0x1A, 0x1A, 0x2E)

CHAPTER_ORDER = [
    "00_document_introduction.md",
    "01_executive_summary.md",
    "02_fraud_landscape.md",
    "02b_case_studies_and_un_data.md",
    "03_ethiopia_market.md",
    "04_solution_overview.md",
    "05_message_flows.md",
    "06_sas_fsm_timeouts.md",
    "07_camara_open_gateway.md",
    "08_gsma_fasg_security.md",
    "09_security_compliance.md",
    "10_commercial_model.md",
    "11_implementation_roadmap.md",
    "12_appendices.md",
]


def set_run_font(run, size=11, bold=False, color=DARK, name="Calibri"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def setup_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Restlink Silent Authentication Proposal  ·  Ethiopia  ·  Page ")
        set_run_font(run, size=9, color=RGBColor(0x6B, 0x72, 0x80))
        add_page_number(p)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for i in range(1, 4):
        try:
            styles[f"Heading {i}"].font.color.rgb = GREEN
            styles[f"Heading {i}"].font.name = "Calibri"
        except KeyError:
            pass
    return doc


def make_charts():
    CHARTS.mkdir(parents=True, exist_ok=True)
    paths = {}

    # Fraud / market chart from fraud_stats.json
    fraud_path = ASSETS / "fraud_stats.json"
    if fraud_path.exists():
        data = json.loads(fraud_path.read_text())
        # support list or {metrics: [...]}
        items = data if isinstance(data, list) else data.get("metrics", data.get("stats", []))
        # pick numeric chartable items
        chartable = []
        for it in items:
            if not isinstance(it, dict):
                continue
            label = it.get("label") or it.get("name")
            val = it.get("value")
            if label is None or val is None:
                continue
            try:
                fval = float(val)
            except (TypeError, ValueError):
                continue
            chartable.append((str(label)[:40], fval, it.get("unit", "")))
        if chartable[:8]:
            fig, ax = plt.subplots(figsize=(9, 4.5))
            labels = [c[0] for c in chartable[:8]]
            vals = [c[1] for c in chartable[:8]]
            colors = ["#078930", "#0D7377", "#C9A227", "#DA121A", "#078930", "#0D7377", "#FCDD09", "#1A1A2E"]
            ax.barh(labels[::-1], vals[::-1], color=colors[: len(vals)][::-1])
            ax.set_xlabel("Value (mixed units — see source table)")
            ax.set_title("Selected cyber / mobile identity indicators")
            fig.tight_layout()
            p = CHARTS / "fraud_indicators.png"
            fig.savefig(p, dpi=150)
            plt.close(fig)
            paths["fraud"] = p

    # ROI chart
    roi_path = ASSETS / "roi_illustrative.json"
    if roi_path.exists():
        roi = json.loads(roi_path.read_text())
        scenarios = roi.get("scenarios") or roi.get("chart_series") or []
        if isinstance(scenarios, dict):
            # chart_series style
            pass
        users, sms, silent, fraud = [], [], [], []
        if isinstance(scenarios, list) and scenarios and "users" in scenarios[0]:
            for s in scenarios:
                users.append(str(s.get("users", s.get("label", ""))))
                sms.append(float(s.get("sms_otp_annual_usd", 0)))
                silent.append(float(s.get("silent_auth_annual_usd", 0)))
                fraud.append(float(s.get("fraud_loss_avoided_usd", 0)))
        elif isinstance(roi.get("chart_series"), dict):
            cs = roi["chart_series"]
            users = [str(x) for x in cs.get("users", ["100k", "500k", "1M"])]
            sms = [float(x) for x in cs.get("sms_otp_annual_usd", [0, 0, 0])]
            silent = [float(x) for x in cs.get("silent_auth_annual_usd", [0, 0, 0])]
            fraud = [float(x) for x in cs.get("fraud_loss_avoided_usd", [0, 0, 0])]
        if users:
            import numpy as np
            x = np.arange(len(users))
            w = 0.25
            fig, ax = plt.subplots(figsize=(8.5, 4.5))
            ax.bar(x - w, sms, w, label="SMS OTP only (USD/yr)", color="#DA121A")
            ax.bar(x, silent, w, label="Blended / Silent Auth (USD/yr)", color="#078930")
            ax.bar(x + w, fraud, w, label="Fraud loss avoided (USD/yr)", color="#C9A227")
            ax.set_xticks(x)
            ax.set_xticklabels(users)
            ax.set_ylabel("USD (illustrative)")
            ax.set_title("ILLUSTRATIVE ROI — Silent Auth vs SMS OTP")
            ax.legend(fontsize=8)
            fig.tight_layout()
            p = CHARTS / "roi_illustrative.png"
            fig.savefig(p, dpi=150)
            plt.close(fig)
            paths["roi"] = p

    # Two-strategy pie-ish bar
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.bar(["Strategy A\nReplace OTP\n(Silent Auth)", "Strategy B\nProtect OTP\n(SMS FW)"],
           [70, 30], color=["#078930", "#0D7377"])
    ax.set_ylabel("Target share of auth events (%)")
    ax.set_title("Target mix after ramp-up (illustrative)")
    ax.set_ylim(0, 100)
    fig.tight_layout()
    p = CHARTS / "strategy_mix.png"
    fig.savefig(p, dpi=150)
    plt.close(fig)
    paths["strategy"] = p

    # Phase roadmap Gantt-like
    fig, ax = plt.subplots(figsize=(8.5, 3.5))
    phases = ["Phase 1: MAP ATI/PSI", "Phase 2: Diameter + SIM Swap", "Phase 3: TS.43 / NV2"]
    starts = [0, 4, 9]
    widths = [4, 5, 6]
    colors = ["#078930", "#0D7377", "#C9A227"]
    ax.barh(phases[::-1], widths[::-1], left=starts[::-1], color=colors[::-1], height=0.5)
    ax.set_xlabel("Months from kickoff")
    ax.set_title("Implementation roadmap (illustrative)")
    fig.tight_layout()
    p = CHARTS / "roadmap.png"
    fig.savefig(p, dpi=150)
    plt.close(fig)
    paths["roadmap"] = p

    return paths


def add_cover(doc: Document):
    for _ in range(2):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RESTLINK")
    set_run_font(r, size=18, bold=True, color=GREEN)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Silent Authentication for Government & Banks")
    set_run_font(r, size=26, bold=True, color=DARK)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Technical & Commercial Proposal — Federal Democratic Republic of Ethiopia")
    set_run_font(r, size=14, color=DARK)

    doc.add_paragraph()
    meta = [
        "Document: Restlink Silent Authentication Proposal v3",
        "Based on: Restlink_Silent_AuthProposal_v3.pptx",
        "Audience: Government digital services · Banks · Ethio Telecom partnership",
        "Classification: Proposal — For Official Use",
        "Date: July 2026",
    ]
    for line in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, size=11, color=DARK)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run(
        "Restlink is a VAS / adapter layer on Ethio Telecom signalling & data. "
        "It does not take SMS interconnect revenue. Fallback OTP remains operator-billed."
    )
    set_run_font(r, size=10, color=RGBColor(0x0D, 0x73, 0x77))
    doc.add_page_break()


def add_toc(doc: Document):
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "0. Document Introduction & Slide Mapping",
        "1. Executive Summary",
        "2. Global Fraud Landscape & Evidence Base",
        "2b. Case Studies & United Nations / ITU Data",
        "3. Ethiopia Market & Stakeholder Context",
        "4. Solution Overview — Two-Stage Silent Auth",
        "5. Protocol Message Flows (ATI / PSI / SAI / Diameter)",
        "6. SAS State Machine, Timeouts & jSS7 Hooks",
        "7. CAMARA & Open Gateway APIs",
        "8. GSMA FASG Signalling & SMS Security",
        "9. Security & Compliance Checklist",
        "10. Commercial Model & Illustrative ROI",
        "11. Implementation Roadmap",
        "12. Appendices — Glossary, References, Document Control",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        for run in p.runs:
            set_run_font(run, size=12)
    doc.add_page_break()


def parse_table_lines(lines: list[str]) -> list[list[str]] | None:
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return rows or None


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            text = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_run_font(run, size=9, bold=(i == 0), color=DARK if i else RGBColor(0xFF, 0xFF, 0xFF))
            if i == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "078930")
                shading.set(qn("w:val"), "clear")
                cell._tc.get_or_add_tcPr().append(shading)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
    doc.add_paragraph()


def add_markdown_file(doc: Document, path: Path, charts: dict):
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    i = 0
    in_code = False
    code_buf = []
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            rows = parse_table_lines(table_buf)
            if rows:
                add_table(doc, rows)
            table_buf = []

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            flush_table()
            if in_code:
                # skip mermaid/code rendering — add caption
                lang = code_buf[0] if code_buf else ""
                p = doc.add_paragraph()
                r = p.add_run(f"[Diagram / code block omitted in DOCX — see markdown source: {path.name}]")
                set_run_font(r, size=9, color=RGBColor(0x6B, 0x72, 0x80))
                code_buf = []
                in_code = False
            else:
                in_code = True
                code_buf = [line.strip()[3:]]
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("|"):
            table_buf.append(line)
            i += 1
            # peek
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                flush_table()
            continue
        else:
            flush_table()

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            for run in p.runs:
                set_run_font(run, size=11)
        elif re.match(r"^\d+\.\s", line.strip()):
            p = doc.add_paragraph(re.sub(r"^\d+\.\s", "", line.strip()), style="List Number")
            for run in p.runs:
                set_run_font(run, size=11)
        elif line.strip() == "":
            pass
        elif line.strip().startswith("---"):
            pass
        else:
            # strip bold markers lightly
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line.strip())
            clean = re.sub(r"`([^`]+)`", r"\1", clean)
            p = doc.add_paragraph(clean)
            for run in p.runs:
                set_run_font(run, size=11)
        i += 1

    flush_table()

    # inject charts after certain chapters
    name = path.name
    if name.startswith("02_") and charts.get("fraud"):
        doc.add_heading("Figure — Selected indicators", level=2)
        doc.add_picture(str(charts["fraud"]), width=Inches(6.2))
        p = doc.add_paragraph("Figure: Mixed-unit indicators from proposal data set (see fraud_stats.json).")
        for run in p.runs:
            set_run_font(run, size=9, color=RGBColor(0x6B, 0x72, 0x80))
    if name.startswith("04_") and charts.get("strategy"):
        doc.add_heading("Figure — Target auth mix", level=2)
        doc.add_picture(str(charts["strategy"]), width=Inches(5.5))
    if name.startswith("10_") and charts.get("roi"):
        doc.add_heading("Figure — Illustrative ROI scenarios", level=2)
        doc.add_picture(str(charts["roi"]), width=Inches(6.2))
        p = doc.add_paragraph("Figure: ILLUSTRATIVE USD figures — not a binding price quote.")
        for run in p.runs:
            set_run_font(run, size=9, color=RGBColor(0x6B, 0x72, 0x80))
    if name.startswith("11_") and charts.get("roadmap"):
        doc.add_heading("Figure — Phased roadmap", level=2)
        doc.add_picture(str(charts["roadmap"]), width=Inches(6.2))


def estimate_pages(doc: Document) -> int:
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    tables = len(doc.tables)
    # Formal proposal: ~280 words/page body + tables ~0.35 page each + cover/toc/charts ~6
    return max(1, int(words / 280 + tables * 0.35 + 6))


def build():
    charts = make_charts()
    doc = setup_doc()
    # Slightly more spacing for page count / readability
    style = doc.styles["Normal"]
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    add_cover(doc)
    add_toc(doc)

    missing = []
    for name in CHAPTER_ORDER:
        path = CHAPTERS / name
        if not path.exists():
            missing.append(name)
            continue
        add_markdown_file(doc, path, charts)
        doc.add_page_break()

    if missing:
        print("WARNING missing chapters:", missing)

    doc.save(OUT)
    est = estimate_pages(doc)
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"Saved → {OUT}")
    print(f"Paragraphs with text: {sum(1 for p in doc.paragraphs if p.text.strip())}")
    print(f"Tables: {len(doc.tables)}")
    print(f"Approx words (body paras): {words}")
    print(f"Estimated pages: ~{est}")
    print(f"Charts: {list(charts)}")
    print(f"File size: {OUT.stat().st_size // 1024} KB")


if __name__ == "__main__":
    build()
